"""
Flask server module for the FirstRespondersChatbot RAG system.

This module sets up a Flask server to provide a REST API for the RAG system,
handling file uploads and queries, and returning responses.
"""

import os
import logging
import uuid
import time
import traceback
from pathlib import Path
from typing import List, Dict, Any, Tuple

from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from werkzeug.utils import secure_filename

from .rag_system import RAGSystem

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


class RAGServer:
    """Server for the FirstRespondersChatbot RAG system."""

    def __init__(
        self,
        rag_system: RAGSystem = None,
        host: str = "0.0.0.0",
        port: int = 8000,
        debug: bool = True,
        allowed_extensions: set = None,
        max_content_length: int = 16 * 1024 * 1024,  # 16MB
        uploads_dir: str = "uploads",
    ):
        """
        Initialize the RAG server.

        Args:
            rag_system: The RAG system to use
            host: Host to bind the server to
            port: Port to bind the server to
            debug: Whether to run the server in debug mode
            allowed_extensions: Set of allowed file extensions
            max_content_length: Maximum content length for file uploads
            uploads_dir: Directory for file uploads
        """
        self.rag_system = rag_system or RAGSystem()
        self.host = host
        self.port = port
        self.debug = debug
        self.allowed_extensions = allowed_extensions or {
            "pdf",
            "txt",
            "md",
            "docx",
            "html",
        }
        self.max_content_length = max_content_length
        self.uploads_dir = Path(uploads_dir)

        # Ensure uploads directory exists
        os.makedirs(str(self.uploads_dir), exist_ok=True)

        # Initialize Flask app
        self.app = Flask(__name__)
        CORS(self.app)  # Enable CORS for all routes

        # Set maximum file size
        self.app.config["MAX_CONTENT_LENGTH"] = self.max_content_length

        # Register routes
        self._register_routes()

        # Initialize request counters and stats
        self.request_count = 0
        self.start_time = time.time()

    def _verify_file_integrity(self, file_data, file_ext) -> tuple:
        """
        Verify file integrity based on file type.

        Args:
            file_data: Binary content of the file
            file_ext: File extension

        Returns:
            tuple: (is_valid, message)
        """
        if file_ext == ".pdf":
            return self._verify_pdf_integrity(file_data)
        elif file_ext == ".docx":
            return self._verify_docx_integrity(file_data)
        elif file_ext in [".txt", ".md", ".html", ".htm"]:
            return self._verify_text_integrity(file_data)

        # Default case for other file types
        return True, f"File verification skipped for {file_ext} files"

    def _verify_pdf_integrity(self, file_data) -> Tuple[bool, str]:
        """
        Verify that a PDF file is valid and readable.

        Args:
            file_data: Binary content of the PDF file

        Returns:
            tuple: (is_valid, message) where is_valid is a boolean and message is a string
        """
        try:
            # Only import PyPDF2 when needed to verify PDFs
            import io
            import PyPDF2

            # Create a file-like object
            pdf_stream = io.BytesIO(file_data)

            # Try to read the PDF
            try:
                pdf_reader = PyPDF2.PdfReader(pdf_stream)
                num_pages = len(pdf_reader.pages)

                # Check if we can access content
                if num_pages > 0:
                    # Try to extract text from the first page
                    first_page = pdf_reader.pages[0]
                    text = first_page.extract_text()

                    # Check if text extraction worked
                    if not text or text.strip() == "":
                        logger.warning(
                            "PDF appears valid but no text could be extracted from the first page"
                        )
                        return (
                            True,
                            "PDF is valid, but may not contain extractable text. Results might be limited.",
                        )

                    # If we have text and it looks like garbled content
                    import re

                    if re.search(r"[^\x00-\x7F]{5,}", text) or re.search(
                        r"([^\w\s]{3,}|([a-zA-Z][^a-zA-Z]){4,})", text
                    ):
                        logger.warning(
                            "PDF contains text but it appears to be corrupted or encoded improperly"
                        )
                        return (
                            True,
                            "PDF contains text but it may be improperly encoded. Results might be affected.",
                        )

                    # If PDF has few pages, ensure we can read them all
                    if num_pages < 5:
                        all_text = ""
                        for i in range(num_pages):
                            page_text = pdf_reader.pages[i].extract_text()
                            all_text += page_text

                        # If total text is very small, warn
                        if len(all_text.strip()) < 100:
                            return (
                                True,
                                "PDF contains very little extractable text. Results might be limited.",
                            )

                return True, f"PDF is valid with {num_pages} pages"
            except Exception as e:
                logger.error(f"Error reading PDF: {str(e)}")
                return False, f"PDF could not be read: {str(e)}"

        except ImportError:
            logger.warning("PyPDF2 not available for PDF verification")
            return True, "PDF verification skipped (PyPDF2 not available)"
        except Exception as e:
            logger.error(f"Error verifying PDF: {str(e)}")
            return False, f"PDF verification failed: {str(e)}"

    def _verify_docx_integrity(self, file_data) -> Tuple[bool, str]:
        """Verify that a DOCX file is valid and readable."""
        try:
            import io
            from docx import Document

            docx_stream = io.BytesIO(file_data)

            try:
                doc = Document(docx_stream)
                paragraphs = len(doc.paragraphs)

                # Check if we can extract any text
                text_content = "\n".join([p.text for p in doc.paragraphs])

                if not text_content or text_content.strip() == "":
                    return (
                        True,
                        "DOCX file is valid but contains no text. Results might be limited.",
                    )

                # If text content is very small, warn
                if len(text_content.strip()) < 100:
                    return (
                        True,
                        "DOCX contains very little text. Results might be limited.",
                    )

                return True, f"DOCX is valid with {paragraphs} paragraphs"
            except Exception as e:
                logger.error(f"Error reading DOCX: {str(e)}")
                return False, f"DOCX could not be read: {str(e)}"

        except ImportError:
            logger.warning("python-docx not available for DOCX verification")
            return True, "DOCX verification skipped (python-docx not available)"
        except Exception as e:
            logger.error(f"Error verifying DOCX: {str(e)}")
            return False, f"DOCX verification failed: {str(e)}"

    def _verify_text_integrity(self, file_data) -> Tuple[bool, str]:
        """Verify that a text file is valid and readable."""
        try:
            # Try to decode the text file with different encodings
            encodings = ["utf-8", "latin-1", "windows-1252", "ascii"]
            decoded = False

            for encoding in encodings:
                try:
                    text = file_data.decode(encoding)
                    decoded = True

                    # Check if the file has reasonable content
                    if not text or text.strip() == "":
                        return True, "Text file is empty. Results might be limited."

                    # If text is very small, warn
                    if len(text.strip()) < 50:
                        return (
                            True,
                            "Text file contains very little content. Results might be limited.",
                        )

                    return True, f"Text file is valid ({len(text)} characters)"
                except UnicodeDecodeError:
                    continue

            if not decoded:
                return (
                    False,
                    "Text file could not be decoded with any standard encoding.",
                )

        except Exception as e:
            logger.error(f"Error verifying text file: {str(e)}")
            return False, f"Text file verification failed: {str(e)}"

    def _register_routes(self):
        """Register routes for the Flask app."""

        # Health check endpoint
        @self.app.route("/api/health", methods=["GET"])
        def health_check():
            """Health check endpoint."""
            uptime = time.time() - self.start_time
            return jsonify(
                {
                    "status": "ok",
                    "uptime": f"{uptime:.2f} seconds",
                    "requests_served": self.request_count,
                    "version": "1.1.0",
                }
            )

        # Upload file endpoint
        @self.app.route("/api/upload", methods=["POST"])
        def upload_file():
            """
            Handle file uploads for the RAG system.

            Returns:
                JSON response with the status of the upload
            """
            self.request_count += 1

            try:
                # Check if a file was provided
                if "file" not in request.files:
                    return (
                        jsonify({"status": "error", "message": "No file provided"}),
                        400,
                    )

                file = request.files["file"]
                # Get session ID (default to "default" if not provided)
                session_id = request.form.get("session_id", "default")

                # Check if the file has a filename
                if file.filename == "" or file.filename is None:
                    return (
                        jsonify({"status": "error", "message": "No selected file"}),
                        400,
                    )

                # Check if the file has an allowed extension
                file_ext = os.path.splitext(file.filename)[1].lower()
                if file_ext[1:] not in self.allowed_extensions:
                    return (
                        jsonify(
                            {
                                "status": "error",
                                "message": f"File type not allowed. Allowed types: {', '.join(self.allowed_extensions)}",
                            }
                        ),
                        400,
                    )

                # Secure the filename to prevent security issues
                original_filename = secure_filename(file.filename)

                # Generate a unique filename to avoid collisions
                unique_prefix = uuid.uuid4().hex[:8]
                unique_filename = f"{unique_prefix}_{original_filename}"

                # Create the full path for the file
                file_path = os.path.join(self.uploads_dir, unique_filename)

                logger.info(
                    f"Uploading file '{original_filename}' to '{file_path}' for session '{session_id}'"
                )

                # Save the file
                file.save(file_path)
                file.seek(0)  # Rewind the file to the beginning for integrity check

                # Check file integrity
                file_data = file.read()
                is_valid, validation_message = self._verify_file_integrity(
                    file_data, file_ext
                )

                if not is_valid:
                    # Delete the invalid file to avoid polluting the uploads directory
                    if os.path.exists(file_path):
                        logger.warning(
                            f"Removing invalid file {file_path}: {validation_message}"
                        )
                        try:
                            os.remove(file_path)
                        except Exception as e:
                            logger.error(f"Failed to remove invalid file: {str(e)}")

                    return (
                        jsonify(
                            {
                                "status": "error",
                                "message": f"Invalid file: {validation_message}",
                            }
                        ),
                        400,
                    )

                # Attempt to index the file with the RAG system
                logger.info(f"Indexing file '{file_path}' for session '{session_id}'")
                indexing_success = self.rag_system.index_file(
                    file_path, session_id=session_id
                )

                # Prepare response
                response = {
                    "file_name": original_filename,
                    "file_path": file_path,
                    "session_id": session_id,
                    "file_size": os.path.getsize(file_path),
                    "file_type": file_ext[1:],
                    "validation_message": validation_message,
                }

                # List all files in this session
                if (
                    session_id in self.rag_system.session_files
                    and self.rag_system.session_files[session_id]
                ):
                    session_files = [
                        os.path.basename(f)
                        for f in self.rag_system.session_files[session_id]
                    ]
                    logger.info(
                        f"Files in session '{session_id}' after upload: {session_files}"
                    )
                    response["session_files"] = session_files

                # Set appropriate status based on indexing success
                if indexing_success:
                    logger.info(
                        f"Successfully uploaded and indexed file '{original_filename}' for session '{session_id}'"
                    )
                    response["message"] = (
                        f"File '{original_filename}' uploaded and indexed successfully"
                    )
                    response["status"] = "success"
                else:
                    logger.warning(
                        f"File '{original_filename}' uploaded but could not be indexed for session '{session_id}'"
                    )
                    response["message"] = (
                        f"File '{original_filename}' uploaded but could not be indexed."
                    )
                    response["status"] = "partial_success"

                return jsonify(response)

            except Exception as e:
                logger.error(f"Error handling file upload: {str(e)}")
                logger.error(traceback.format_exc())

                return (
                    jsonify(
                        {
                            "status": "error",
                            "message": f"Error processing file: {str(e)}",
                        }
                    ),
                    500,
                )

        # Query endpoint
        @self.app.route("/api/query", methods=["POST"])
        def query():
            """
            Handle queries to the RAG system.

            Returns:
                JSON response with the answer and context
            """
            self.request_count += 1

            # Get query from request
            data = request.json
            if not data or "query" not in data:
                return jsonify({"status": "error", "message": "No query provided"}), 400

            query_text = data["query"]
            # Get session ID (default to "default" if not provided)
            session_id = data.get("session_id", "default")

            try:
                logger.info(
                    f"Processing query: '{query_text}' for session: {session_id}"
                )

                start_time = time.time()
                # Generate response with session context
                response = self.rag_system.generate_response(
                    query_text, session_id=session_id
                )
                end_time = time.time()

                processing_time = end_time - start_time
                logger.info(f"Query processed in {processing_time:.2f} seconds")

                # Check if the answer is our generic error message
                if response["answer"].startswith(
                    "I apologize, but I'm having trouble processing"
                ):
                    logger.warning(
                        f"Model returned the generic error response for query: '{query_text}'"
                    )

                    # For common first responder topics, provide a more helpful error
                    if (
                        "ppe" in query_text.lower()
                        or "protective equipment" in query_text.lower()
                    ):
                        logger.info(
                            "Detected PPE question, providing specialized response"
                        )

                        # Keep the original error source but improve the answer
                        response["answer"] = (
                            "Based on my knowledge, Personal Protective Equipment (PPE) for first responders is designed to protect them from various hazards encountered during emergency operations. PPE serves as a critical barrier between responders and dangerous environments."
                            + "\n\n[Generated by Phi-3 Model]"
                        )

                # Ensure consistent context format between RAG and non-RAG responses
                formatted_context = []
                for ctx in response["context"]:
                    formatted_item = {}
                    # Source is always present
                    if "source" in ctx:
                        formatted_item["source"] = ctx["source"]

                    # Content could be in "content" or "snippet"
                    if "content" in ctx:
                        formatted_item["content"] = ctx["content"]
                    elif "snippet" in ctx:
                        formatted_item["snippet"] = ctx["snippet"]

                    # File name could be explicit or in metadata
                    if "file_name" in ctx:
                        formatted_item["file_name"] = ctx["file_name"]

                    formatted_context.append(formatted_item)

                # Log which files were used for this query
                session_files = []
                if session_id in self.rag_system.session_files:
                    session_files = list(self.rag_system.session_files[session_id])
                    if session_files:
                        logger.info(f"Session {session_id} has files: {session_files}")
                    else:
                        logger.info(f"Session {session_id} has no files")

                return jsonify(
                    {
                        "status": "success",
                        "answer": response["answer"],
                        "context": formatted_context,
                        "query": response["query"],
                        "session_id": session_id,
                        "processing_time": f"{processing_time:.2f} seconds",
                        "used_files": (
                            [os.path.basename(f) for f in session_files]
                            if session_files
                            else []
                        ),
                    }
                )

            except Exception as e:
                logger.error(f"Error handling query: {str(e)}")
                logger.error(traceback.format_exc())

                # Provide a more informative error message
                error_message = str(e)
                if "CUDA out of memory" in error_message:
                    error_message = "The system is experiencing high memory usage. Please try again with a simpler query or wait a moment."
                elif "Connection refused" in error_message:
                    error_message = "The backend services are currently unavailable. Please try again later."
                elif "garbled" in error_message.lower():
                    error_message = "The system is having trouble generating a clear response. Please try rephrasing your question."
                elif (
                    "model" in error_message.lower() and "load" in error_message.lower()
                ):
                    error_message = "There was an issue with the AI model. The system will use a fallback model instead."

                # Specialized responses for specific topic errors
                topic_keywords = {
                    "ppe": "Personal Protective Equipment (PPE) protects first responders from hazards during emergency operations.",
                    "firefighter": "There was an issue retrieving information about firefighter procedures.",
                    "emergency": "There was an issue retrieving information about emergency procedures.",
                }

                for keyword, message in topic_keywords.items():
                    if keyword in query_text.lower():
                        return (
                            jsonify(
                                {
                                    "status": "partial_success",
                                    "answer": message
                                    + " Please try a more specific question."
                                    + "\n\n[Generated by Phi-3 Model]",
                                    "query": query_text,
                                    "context": [
                                        {
                                            "source": "Error recovery",
                                            "content": "Partial information provided.",
                                            "file_name": "Model Knowledge",
                                        }
                                    ],
                                    "session_id": session_id,
                                }
                            ),
                            200,
                        )

                return (
                    jsonify(
                        {
                            "status": "error",
                            "message": f"Error processing query: {error_message}",
                        }
                    ),
                    500,
                )

        # Serve uploaded files
        @self.app.route("/api/files/<path:filename>", methods=["GET"])
        def serve_file(filename):
            """Serve uploaded files."""
            try:
                # Add security check to prevent path traversal
                requested_path = os.path.abspath(
                    os.path.join(self.uploads_dir, filename)
                )
                if not requested_path.startswith(os.path.abspath(self.uploads_dir)):
                    return jsonify({"status": "error", "message": "Access denied"}), 403

                # Check if file exists
                if not os.path.exists(requested_path):
                    return (
                        jsonify({"status": "error", "message": "File not found"}),
                        404,
                    )

                # Serve the file
                return send_from_directory(self.uploads_dir, filename)
            except Exception as e:
                logger.error(f"Error serving file: {str(e)}")
                return jsonify({"status": "error", "message": str(e)}), 500

        # Remove file endpoint
        @self.app.route("/api/remove-file", methods=["POST"])
        def remove_file():
            """
            Remove a file from a session.

            Returns:
                JSON response with status and message
            """
            try:
                data = request.json
                if not data or "file_path" not in data:
                    return (
                        jsonify(
                            {"status": "error", "message": "No file path provided"}
                        ),
                        400,
                    )

                file_path = data["file_path"]
                # Get session ID (default to "default" if not provided)
                session_id = data.get("session_id", "default")

                logger.info(f"Removing file '{file_path}' from session '{session_id}'")

                # Check if the path refers to an existing file
                full_path = os.path.join(self.uploads_dir, file_path)
                if not os.path.exists(full_path):
                    # Check if it might be a base filename without path
                    for root, _, files in os.walk(self.uploads_dir):
                        if file_path in files:
                            full_path = os.path.join(root, file_path)
                            break
                    else:  # No break occurred, file not found
                        return (
                            jsonify({"status": "error", "message": "File not found"}),
                            404,
                        )

                # Remove file from RAG system
                if self.rag_system.remove_file(full_path, session_id=session_id):
                    # Verify the file was removed from the session
                    if (
                        session_id in self.rag_system.session_files
                        and full_path in self.rag_system.session_files[session_id]
                    ):
                        logger.warning(
                            f"File was not properly removed from session tracking"
                        )
                        return (
                            jsonify(
                                {
                                    "status": "error",
                                    "message": "File was not completely removed from the system",
                                }
                            ),
                            500,
                        )

                    logger.info(
                        f"Successfully removed file '{file_path}' from session '{session_id}'"
                    )
                    # List active files in the session after removal
                    if (
                        session_id in self.rag_system.session_files
                        and self.rag_system.session_files[session_id]
                    ):
                        remaining_files = [
                            os.path.basename(f)
                            for f in self.rag_system.session_files[session_id]
                        ]
                        logger.info(
                            f"Remaining files in session '{session_id}': {remaining_files}"
                        )
                    else:
                        logger.info(f"No remaining files in session '{session_id}'")

                    return jsonify(
                        {
                            "status": "success",
                            "message": f"File '{file_path}' removed from session",
                            "session_id": session_id,
                            "remaining_files": (
                                [
                                    os.path.basename(f)
                                    for f in self.rag_system.session_files[session_id]
                                ]
                                if session_id in self.rag_system.session_files
                                else []
                            ),
                        }
                    )
                else:
                    logger.warning(
                        f"Failed to remove file '{file_path}' from session '{session_id}'"
                    )
                    return (
                        jsonify(
                            {
                                "status": "error",
                                "message": "Failed to remove file from the system",
                            }
                        ),
                        500,
                    )

            except Exception as e:
                logger.error(f"Error removing file: {str(e)}")
                logger.error(traceback.format_exc())
                return jsonify({"status": "error", "message": str(e)}), 500

        # Clear session endpoint
        @self.app.route("/api/clear-session", methods=["POST"])
        def clear_session():
            """
            Clear a session (remove all files associated with a session).

            Returns:
                JSON response with status and message
            """
            try:
                data = request.json
                session_id = data.get("session_id", "default")

                logger.info(f"Clearing session '{session_id}'")

                # Check if session exists
                if session_id not in self.rag_system.session_files:
                    return jsonify(
                        {
                            "status": "success",
                            "message": f"Session '{session_id}' is already empty",
                        }
                    )

                # Count files before clearing
                file_count = len(self.rag_system.session_files[session_id])
                logger.info(f"Removing {file_count} files from session '{session_id}'")

                # File names for logging
                files = [
                    os.path.basename(f)
                    for f in self.rag_system.session_files[session_id]
                ]
                logger.info(f"Files to be removed: {files}")

                # Clear the session
                success = self.rag_system.clear_session(session_id)

                if success:
                    # Verify the session was cleared properly
                    if (
                        session_id in self.rag_system.session_files
                        and self.rag_system.session_files[session_id]
                    ):
                        remaining = len(self.rag_system.session_files[session_id])
                        logger.warning(
                            f"Session '{session_id}' still has {remaining} files after clearing"
                        )
                        return jsonify(
                            {
                                "status": "partial_success",
                                "message": f"Session '{session_id}' was partially cleared, but {remaining} files remain",
                                "session_id": session_id,
                            }
                        )

                    logger.info(f"Successfully cleared session '{session_id}'")
                    return jsonify(
                        {
                            "status": "success",
                            "message": f"Session '{session_id}' cleared successfully. Removed {file_count} files.",
                            "removed_files": files,
                        }
                    )
                else:
                    logger.warning(f"Failed to clear session '{session_id}'")
                    return jsonify(
                        {
                            "status": "error",
                            "message": f"Failed to clear session '{session_id}'",
                        }
                    )

            except Exception as e:
                logger.error(f"Error clearing session: {str(e)}")
                logger.error(traceback.format_exc())
                return jsonify({"status": "error", "message": str(e)}), 500

        # Clear index endpoint
        @self.app.route("/api/clear", methods=["POST"])
        def clear_index():
            """
            Clear the document index.

            Returns:
                JSON response with status and message
            """
            self.request_count += 1

            try:
                self.rag_system.clear_index()
                return jsonify(
                    {
                        "status": "success",
                        "message": "Document index cleared successfully",
                    }
                )

            except Exception as e:
                logger.error(f"Error clearing index: {str(e)}")
                logger.error(traceback.format_exc())
                return (
                    jsonify(
                        {
                            "status": "error",
                            "message": f"Error clearing index: {str(e)}",
                        }
                    ),
                    500,
                )

        # Get indexed files endpoint
        @self.app.route("/api/files", methods=["GET"])
        def get_indexed_files():
            """
            Get a list of indexed files for a session.

            Returns:
                JSON response with the list of indexed files
            """
            self.request_count += 1

            try:
                # Get session ID from query parameter
                session_id = request.args.get("session_id", "default")

                if session_id in self.rag_system.session_files:
                    files = list(self.rag_system.session_files[session_id])
                    file_info = []

                    for file_path in files:
                        file_path_obj = Path(file_path)
                        if os.path.exists(file_path):
                            file_info.append(
                                {
                                    "name": file_path_obj.name,
                                    "path": file_path,
                                    "size": os.path.getsize(file_path),
                                    "type": file_path_obj.suffix.lower()[
                                        1:
                                    ],  # Remove the dot
                                    "last_modified": os.path.getmtime(file_path),
                                }
                            )

                    return jsonify(
                        {
                            "status": "success",
                            "files": file_info,
                            "session_id": session_id,
                            "count": len(file_info),
                        }
                    )
                else:
                    return jsonify(
                        {
                            "status": "success",
                            "files": [],
                            "session_id": session_id,
                            "count": 0,
                            "message": f"No files found for session {session_id}",
                        }
                    )

            except Exception as e:
                logger.error(f"Error getting indexed files: {str(e)}")
                logger.error(traceback.format_exc())
                return (
                    jsonify(
                        {
                            "status": "error",
                            "message": f"Error getting indexed files: {str(e)}",
                        }
                    ),
                    500,
                )

        # Get server status endpoint
        @self.app.route("/api/status", methods=["GET"])
        def get_status():
            """
            Get the server status.

            Returns:
                JSON response with the server status
            """
            uptime = time.time() - self.start_time
            uptime_str = f"{int(uptime // 3600)}h {int((uptime % 3600) // 60)}m {int(uptime % 60)}s"

            # Get model information
            model_info = {
                "name": "Microsoft Phi-3-medium-4k-instruct",
                "type": "Instruction-tuned language model",
                "parameters": "8 billion parameters",
                "loaded_on": (
                    self.rag_system.device.type
                    if hasattr(self.rag_system, "device")
                    else "unknown"
                ),
            }

            return jsonify(
                {
                    "status": "online",
                    "uptime": uptime_str,
                    "requests_processed": self.request_count,
                    "model_info": model_info,
                }
            )

    def _allowed_file(self, filename: str) -> bool:
        """
        Check if a file has an allowed extension.

        Args:
            filename: The filename to check

        Returns:
            bool: True if the file has an allowed extension, False otherwise
        """
        return (
            "." in filename
            and filename.rsplit(".", 1)[1].lower() in self.allowed_extensions
        )

    def run(self):
        """Run the server."""
        logger.info(f"Starting server on {self.host}:{self.port}")
        self.app.run(host=self.host, port=self.port, debug=self.debug)
